"""合同文件解析:docx/python-docx 标题分章,pdf/pypdf 文本层,无文本层抛 NeedsOcr(设计 §4.1)。

统一输出 Document{chapters[], total_chars, source_type: docx|pdf}。
大小校验:解析前文件 ≤2MB(max_bytes),解析后正文总字数 ≤5 万字(max_chars),超限报 ContractTooLongError。
"""
from __future__ import annotations

from pathlib import Path

from pydantic import BaseModel

from agents.contract_review_agent.utils.chapterizer import Chapter, build_chapters


class Document(BaseModel):
    chapters: list[Chapter]
    total_chars: int
    source_type: str


class ContractTooLongError(ValueError):
    pass


class NeedsOcrError(ValueError):
    pass


class UnsupportedTypeError(ValueError):
    pass


def parse_document(path: str | Path, max_bytes: int = 2 * 1024 * 1024,
                   max_chars: int = 50_000) -> Document:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix not in (".docx", ".pdf"):
        raise UnsupportedTypeError(f"不支持的文件类型: {suffix}")
    if not p.exists():
        raise FileNotFoundError(str(p))
    if p.stat().st_size > max_bytes:
        raise ContractTooLongError("文件超过 2MB 限制")
    if suffix == ".docx":
        blocks = _parse_docx(p)
        source_type = "docx"
    else:  # suffix == ".pdf"(前置已拦截非 docx/pdf)
        blocks = _parse_pdf(p)
        source_type = "pdf"
    chapters = build_chapters(blocks)
    total = sum(len(c.text) for c in chapters)
    if total > max_chars:
        raise ContractTooLongError(f"正文超过 5 万字限制(实际 {total})")
    return Document(chapters=chapters, total_chars=total, source_type=source_type)


def _parse_docx(p: Path) -> list[tuple[str, int]]:
    from docx import Document as _Docx

    doc = _Docx(str(p))
    blocks: list[tuple[str, int]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        level = 1 if ("heading" in style or "标题" in style) else 0
        blocks.append((text, level))
    return blocks


def _parse_pdf(p: Path) -> list[tuple[str, int]]:
    from pypdf import PdfReader

    reader = PdfReader(str(p))
    pages: list[str] = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        pages.append(t)
    if not any(pages):
        raise NeedsOcrError("PDF 无文本层,需要 OCR")
    blocks: list[tuple[str, int]] = []
    for page_text in pages:
        for line in page_text.splitlines():
            line = line.strip()
            if not line:
                continue
            blocks.append((line, 1 if _looks_like_heading(line) else 0))
    return blocks


def _looks_like_heading(line: str) -> bool:
    return (line.startswith(("第", "一、", "二、", "三、", "四、", "五、",
                             "六、", "七、", "八、", "九、", "十、"))
            and any(k in line for k in ("章", "节", "条"))) or len(line) <= 30
